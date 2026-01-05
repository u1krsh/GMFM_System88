
import unittest
import tempfile
import sys
import os
from pathlib import Path
from datetime import date
from docx import Document

# Fix import path
PROJECT_ROOT = Path(__file__).resolve().parents[1]
SRC_PATH = PROJECT_ROOT / "src"
if str(SRC_PATH) not in sys.path:
    sys.path.insert(0, str(SRC_PATH))

from gmfm_app.services.docx_import_service import parse_docx, import_assessment_to_db, ImportedAssessment
from gmfm_app.data.database import DatabaseContext, init_db
import sqlite3

class TestDocxImport(unittest.TestCase):
    def setUp(self):
        # Create a temporary filled DOCX file
        self.temp_dir = tempfile.TemporaryDirectory()
        self.docx_path = Path(self.temp_dir.name) / "test_patient.docx"
        self._create_sample_docx(self.docx_path)
        
        # Setup in-memory DB
        self.db_path = Path(self.temp_dir.name) / "test_db.sqlite"
        self.db_context = DatabaseContext(str(self.db_path))
        init_db(self.db_path)

    def tearDown(self):
        self.temp_dir.cleanup()

    def _create_sample_docx(self, path):
        doc = Document()
        doc.add_paragraph("Gross Motor Function Measure Score Sheet")
        doc.add_paragraph("Name: John Doe")
        doc.add_paragraph("Assessment Date: 2023-10-25")
        doc.add_paragraph("Evaluator's name: Dr. Smith")
        
        # Add a table with some scores
        table = doc.add_table(rows=1, cols=3)
        # Header row (though parser doesn't rely on headers, just row content)
        
        # Add some score rows
        # Format: Item Num, Description, Score
        items = [
            ("1.", "Lying Description", "3"),
            ("2.", "Rolling Description", "2"),
            ("10.", "Sitting Description", "1"),
            ("88.", "Jumping Description", "0"),
            ("50.", "Nt Item", "NT"), # Should be ignored
        ]
        
        for num, desc, score in items:
            row = table.add_row()
            row.cells[0].text = num
            row.cells[1].text = desc
            row.cells[2].text = score
            
        doc.save(path)

    def test_parse_docx(self):
        assessment = parse_docx(self.docx_path)
        
        self.assertEqual(assessment.patient_name, "John Doe")
        self.assertEqual(assessment.given_name, "John")
        self.assertEqual(assessment.family_name, "Doe")
        self.assertEqual(assessment.assessment_date, date(2023, 10, 25))
        self.assertEqual(assessment.evaluator_name, "Dr. Smith")
        
        # Check scores
        self.assertEqual(assessment.raw_scores[1], 3)
        self.assertEqual(assessment.raw_scores[2], 2)
        self.assertEqual(assessment.raw_scores[10], 1)
        self.assertEqual(assessment.raw_scores[88], 0)
        self.assertNotIn(50, assessment.raw_scores) # NT should be skipped
        
        self.assertEqual(len(assessment.raw_scores), 4)

    def test_import_to_db(self):
        assessment = parse_docx(self.docx_path)
        patient_id, session_id = import_assessment_to_db(assessment, self.db_context)
        
        with self.db_context.connect() as conn:
            cursor = conn.cursor()
            
            # Verify Patient
            cursor.execute("SELECT given_name, family_name FROM patients WHERE id = ?", (patient_id,))
            p_row = cursor.fetchone()
            self.assertEqual(p_row[0], "John")
            self.assertEqual(p_row[1], "Doe")
            
            # Verify Session
            cursor.execute("SELECT scale, total_score, notes, raw_scores FROM sessions WHERE id = ?", (session_id,))
            s_row = cursor.fetchone()
            self.assertEqual(s_row[0], "88")
            self.assertIn("Evaluator: Dr. Smith", s_row[2])
            
            # 4 items * 3 max = 12 total max.
            # Scores: 3+2+1+0 = 6.
            # 6 / 12 * 100 = 50.0%
            self.assertEqual(s_row[1], 50.0)

if __name__ == "__main__":
    unittest.main()
